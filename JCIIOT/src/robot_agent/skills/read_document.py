"""Read Document Skill — extract a DOCX and build team-owned SOP knowledge.

The competition ships reference SOP markdown under ``knowledge/``.  Those
files are locked.  This skill therefore writes generated markdown only under
``team_submission/knowledge/``, which the planner already loads after the
locked knowledge block.

The DOCX supplies task semantics, not world coordinates.  Coordinates are
resolved by joining the current runtime environment to ``task_config.json``
and then reading the matching station records from the generated semantic
map.  Any missing or ambiguous join fails closed instead of guessing.
"""

from __future__ import annotations

import json
import logging
import re
from pathlib import Path
from typing import Any

from robot_agent.core.types import ExecutionContext, SkillResult
from robot_agent.skills.base import BaseSkill

logger = logging.getLogger(__name__)

_PROJECT_ROOT = Path(__file__).resolve().parents[3]
_TEAM_KNOWLEDGE_ROOT = _PROJECT_ROOT / "team_submission" / "knowledge"
_GENERATED_SOPS_ROOT = _PROJECT_ROOT / "team_submission" / "generated_sops"
_ACTIVE_SOP_PATH = _TEAM_KNOWLEDGE_ROOT / "current_generated_sop.md"
_TASK_CONFIG_PATH = _PROJECT_ROOT / "knowledge" / "task_config.json"
_MAPS_ROOT = (
    _PROJECT_ROOT
    / "robosuite"
    / "robosuite"
    / "environments"
    / "factory_sorting"
    / "maps"
)
_GENERATED_MAPS_ROOT = (
    _PROJECT_ROOT
    / "robosuite"
    / "robosuite"
    / "environments"
    / "factory_sorting"
    / "generated_maps"
)
_NUMBER_WORDS = {
    "one": 1,
    "two": 2,
    "three": 3,
    "four": 4,
    "five": 5,
    "six": 6,
    "seven": 7,
    "eight": 8,
    "nine": 9,
    "ten": 10,
}


def _compact(text: str) -> str:
    return " ".join(str(text).replace("\xa0", " ").split())


def _extract_prompt_lines(paragraphs: list[str]) -> list[str]:
    """Return only the current-task prompt at the start of an official DOCX.

    Several competition documents contain stale task examples in the generic
    SOP body.  The prompt block precedes the first ``JCIIOT 2026 ...`` heading,
    so it is intentionally separated before any task fields are parsed.
    """

    prompt_lines: list[str] = []
    for paragraph in paragraphs:
        line = _compact(paragraph)
        if not line:
            continue
        if re.match(
            r"^JCIIOT\s+2026\s+(?:Standard\s+Operating\s+Procedure|Operation\s+Instruction)",
            line,
            flags=re.IGNORECASE,
        ):
            break
        line = re.sub(r"^Prompt\s*[:：]\s*", "", line, flags=re.IGNORECASE)
        if line:
            prompt_lines.append(line)
        if len(prompt_lines) >= 12:
            break
    return prompt_lines


def _case_number_from_path(path: Path) -> int | None:
    match = re.search(r"\bcase[\s_-]*(\d+)\b", path.stem, flags=re.IGNORECASE)
    if match:
        return int(match.group(1))
    return None


def _case_number_from_task(task: dict[str, Any]) -> int | None:
    """Read a case number from task data instead of a code-side lookup table."""
    for field in ("case_number", "case", "scene_number"):
        value = task.get(field)
        if value is not None:
            try:
                return int(value)
            except (TypeError, ValueError):
                pass

    env_name = str(task.get("env_name") or "")
    match = re.search(r"FactorySorting(\d+)(?:_|$)", env_name, flags=re.IGNORECASE)
    return int(match.group(1)) if match else None


def _parse_prompt_fields(prompt_lines: list[str]) -> dict[str, Any]:
    """Extract human-readable task fields while preserving the raw prompt."""

    text = " ".join(prompt_lines)

    def first(patterns: tuple[str, ...]) -> str:
        for pattern in patterns:
            match = re.search(pattern, text, flags=re.IGNORECASE)
            if match:
                return _compact(match.group(1)).strip(" .,:;\"'")
        return ""

    material = first(
        (
            r"Material\s+Name\s*:\s*(.+?)(?=\s+(?:Starting\s+Location|Target\s+Location|Quantity\s+to\s+Transport)\s*:|$)",
            r"object\s+to\s+be\s+handled\s+is\s+(.+?)\.",
            r"\bobject\s+is\s+(.+?)\.",
            r"need\s+to\s+transport\s+(?:a|an|the)\s+(.+?)\.",
            r"\bMove\s+the\s+(?:one|two|three|four|five|\d+)\s+(.+?)\s+from\s+Pick\s+Station",
        )
    )
    source_label = first(
        (
            r"Starting\s+Location\s*:\s*(Pick\s+Station\s+\d+)",
            r"starting\s+point\s+[\"'](Pick\s+Station\s+\d+)[\"']",
            r"Pick\s+Station\s+is(?:\s+designated\s+as)?\s+(Pick\s+Station\s+\d+)",
            r"\bfrom\s+(Pick\s+Station\s+\d+)",
        )
    )
    target_label = first(
        (
            r"Target\s+Location\s*:\s*(Place\s+Station\s+\d+)",
            r"destination\s+[\"'](Place\s+Station\s+\d+)[\"']",
            r"Place\s+Station\s+is(?:\s+designated\s+as)?\s+(Place\s+Station\s+\d+)",
            r"\bto\s+(Place\s+Station\s+\d+)",
        )
    )

    quantity = 1
    quantity_match = re.search(
        r"Quantity\s+to\s+Transport\s*:\s*(\d+)",
        text,
        flags=re.IGNORECASE,
    )
    if quantity_match:
        quantity = max(1, int(quantity_match.group(1)))
    else:
        quantity_match = re.search(
            r"\bMove\s+the\s+(one|two|three|four|five|six|seven|eight|nine|ten|\d+)\b",
            text,
            flags=re.IGNORECASE,
        )
        if quantity_match:
            raw_quantity = quantity_match.group(1).lower()
            quantity = int(raw_quantity) if raw_quantity.isdigit() else _NUMBER_WORDS[raw_quantity]

    return {
        "prompt_lines": prompt_lines,
        "material": material or "Not deterministically parsed; use the source prompt",
        "source_label": source_label or "Not deterministically parsed; use the source prompt",
        "target_label": target_label or "Not deterministically parsed; use the source prompt",
        "quantity": quantity,
    }


def _resolve_task_entry(
    path: Path,
    *,
    requested_level: str = "",
    requested_environment: str = "",
) -> tuple[dict[str, Any], str, int | None]:
    """Resolve exactly one task from runtime data; never guess or hardcode it."""
    if not _TASK_CONFIG_PATH.exists():
        raise FileNotFoundError(f"Task catalog not found: {_TASK_CONFIG_PATH}")

    task_config = json.loads(_TASK_CONFIG_PATH.read_text(encoding="utf-8"))
    tasks = [
        item
        for item in task_config.get("tasks", [])
        if isinstance(item, dict)
    ]
    if not tasks:
        raise RuntimeError(f"No tasks are defined in {_TASK_CONFIG_PATH}")

    requested_level = _compact(requested_level).upper()
    requested_environment = _compact(requested_environment)
    case_number = _case_number_from_path(path)

    candidates = tasks
    selectors: list[str] = []
    if requested_environment:
        candidates = [
            task
            for task in candidates
            if str(task.get("env_name") or "") == requested_environment
        ]
        selectors.append(f"environment={requested_environment!r}")
    if requested_level:
        candidates = [
            task
            for task in candidates
            if str(task.get("level") or "").upper() == requested_level
        ]
        selectors.append(f"level={requested_level!r}")
    if case_number is not None:
        candidates = [
            task
            for task in candidates
            if _case_number_from_task(task) == case_number
        ]
        selectors.append(f"document case={case_number}")

    if len(candidates) != 1:
        selector_text = ", ".join(selectors) or "no explicit selector"
        matches = [
            {
                "level": task.get("level"),
                "env_name": task.get("env_name"),
                "case_number": _case_number_from_task(task),
            }
            for task in candidates
        ]
        raise RuntimeError(
            "Document could not be mapped to exactly one runtime task "
            f"using {selector_text}. Matches: {matches}. "
            "Provide inputs.level or inputs.environment; coordinates will not be guessed."
        )

    task = candidates[0]
    level = str(task.get("level") or "").upper()
    resolved_case = _case_number_from_task(task)
    return task, level, resolved_case


def _load_scene_map(task: dict[str, Any] | None) -> tuple[dict[str, Any], Path | None]:
    if not task:
        return {}, None
    scene_prefix = str(task.get("scene_prefix") or "")
    candidates = (
        _GENERATED_MAPS_ROOT
        / f"{scene_prefix}_scene_regenerated_semantic_map.json",
        _MAPS_ROOT / f"{scene_prefix}_scene_regenerated.json",
    )
    for map_path in candidates:
        if map_path.exists():
            return json.loads(map_path.read_text(encoding="utf-8")), map_path
    return {}, candidates[0]


def _station_record(scene_map: dict[str, Any], station_name: str) -> dict[str, Any]:
    for group_name in ("input_ports", "output_ports"):
        group = scene_map.get(group_name, {})
        if isinstance(group, dict):
            station = group.get(station_name)
            if isinstance(station, dict):
                return station
    for item in scene_map.get("objects", []):
        if item.get("name") == station_name:
            return item
    return {}


def _format_xy(value: Any) -> str:
    if not isinstance(value, (list, tuple)) or len(value) < 2:
        return "not available"
    return f"({float(value[0]):.6g}, {float(value[1]):.6g})"


def _safe_output_path(level: str, case_number: int | None, inputs: dict[str, Any]) -> Path:
    root = _GENERATED_SOPS_ROOT.resolve()
    requested_dir = inputs.get("sop_output_dir")
    if requested_dir:
        candidate_dir = Path(str(requested_dir))
        if not candidate_dir.is_absolute():
            candidate_dir = _PROJECT_ROOT / candidate_dir
        output_dir = candidate_dir.resolve()
    else:
        output_dir = root

    if output_dir != root and root not in output_dir.parents:
        raise ValueError(
            "Generated SOP archive must stay inside team_submission/generated_sops"
        )

    requested_name = _compact(str(inputs.get("sop_output_name") or ""))
    if requested_name:
        if Path(requested_name).name != requested_name:
            raise ValueError("sop_output_name must be a file name, not a path")
        safe_stem = re.sub(r"[^A-Za-z0-9_.-]+", "_", Path(requested_name).stem)
        file_name = f"{safe_stem}.md"
    elif level:
        file_name = f"generated_sop_{level.lower()}.md"
    elif case_number is not None:
        file_name = f"generated_sop_case_{case_number}.md"
    else:
        file_name = "generated_sop_document.md"

    output_dir.mkdir(parents=True, exist_ok=True)
    output_path = (output_dir / file_name).resolve()
    if output_path.parent != output_dir:
        raise ValueError("Unsafe generated SOP path")
    return output_path


def _build_sop_markdown(
    *,
    source_path: Path,
    level: str,
    case_number: int | None,
    prompt: dict[str, Any],
    task: dict[str, Any] | None,
    scene_map: dict[str, Any],
    map_path: Path | None,
    image_count: int,
    image_descriptions: dict[str, str],
) -> str:
    task = task or {}
    source = str(task.get("source") or "not available")
    target = str(task.get("target") or "not available")
    object_names = [str(item) for item in task.get("object", []) if str(item)]
    quantity = int(prompt.get("quantity") or 1)
    selected_objects = object_names[:quantity] if quantity > 1 else object_names[:1]

    source_station = _station_record(scene_map, source)
    target_station = _station_record(scene_map, target)
    for prompt_field in ("material", "source_label", "target_label"):
        value = str(prompt.get(prompt_field) or "")
        if not value or value.startswith("Not deterministically parsed"):
            raise RuntimeError(
                f"Word prompt field {prompt_field!r} could not be resolved; "
                "SOP generation stopped instead of guessing"
            )
    for station_name, station in ((source, source_station), (target, target_station)):
        for coordinate_field in ("center", "approach"):
            value = station.get(coordinate_field)
            if not isinstance(value, (list, tuple)) or len(value) < 2:
                raise RuntimeError(
                    f"Scene map has no valid {coordinate_field!r} coordinate "
                    f"for station {station_name!r}; SOP generation stopped"
                )

    title_level = level or (f"Case {case_number}" if case_number is not None else "Task")
    lines = [
        f"# {title_level} Team-Generated Task SOP",
        "",
        "<!-- TEAM GENERATED — SAFE TO REGENERATE -->",
        "",
        "## Planner Quick Facts",
        "",
        f"- CURRENT LEVEL: `{level or 'not available'}`",
        f"- MATERIAL: {prompt.get('material')}",
        f"- SOURCE: `{source}` ({prompt.get('source_label')})",
        f"- TARGET: `{target}` ({prompt.get('target_label')})",
        f"- QUANTITY: `{quantity}`",
        "- EXACT OBJECTS: "
        + (", ".join(f"`{name}`" for name in object_names) if object_names else "unresolved; do not guess"),
        f"- REQUIRED PLAN: `move {source} → pick_up → move {target} → place_down`",
        "- The current task prompt and runtime mapping override stale examples in the generic Word SOP body.",
        "",
        f"- Source DOCX: `{source_path.name}`",
        "- Generator: `src/robot_agent/skills/read_document.py`",
        "- Destination: team-owned SOP archive; locked competition knowledge is not modified",
        "",
        "## Authority Rules",
        "",
        "1. The Current Task Prompt below defines the requested material, human station names, and quantity.",
        "2. Runtime environment, station IDs, and object names come from locked `knowledge/task_config.json`.",
        "3. Station geometry comes only from the matching generated semantic map (with scene JSON fallback).",
        "4. Ignore stale task names or station examples that appear later in the generic Word SOP body.",
        "5. Never invent a coordinate, station ID, or object name when a source does not provide it.",
        "",
        "## Current Task Prompt (authoritative)",
        "",
    ]
    prompt_lines = prompt.get("prompt_lines") or ["No prompt block was extracted."]
    lines.extend(f"> {line}" for line in prompt_lines)
    lines.extend(
        [
            "",
            "## Parsed Task Summary",
            "",
            f"- Level: `{level or 'not available'}`",
            f"- Case number: `{case_number if case_number is not None else 'not available'}`",
            f"- Material: {prompt.get('material')}",
            f"- Human source label: {prompt.get('source_label')}",
            f"- Human target label: {prompt.get('target_label')}",
            f"- Quantity: `{quantity}`",
            "",
            "## Runtime Mapping",
            "",
            f"- Scene prefix: `{task.get('scene_prefix', 'not available')}`",
            f"- Environment: `{task.get('env_name', 'not available')}`",
            f"- Source station ID: `{source}`",
            f"- Target station ID: `{target}`",
            f"- Maximum score: `{task.get('max_score', 'not available')}`",
        ]
    )
    if object_names:
        lines.append("- Exact object names, in configured order:")
        lines.extend(f"  - `{name}`" for name in object_names)
    else:
        lines.append("- Exact object names: not available")

    lines.extend(
        [
            "",
            "## Station Geometry",
            "",
            f"- Scene map: `{map_path.name if map_path and map_path.exists() else 'not available'}`",
            f"- Coordinate frame: `{scene_map.get('coordinate_frame', 'not available')}`",
            f"- Source center: `{_format_xy(source_station.get('center'))}`",
            f"- Source navigation approach: `{_format_xy(source_station.get('approach'))}`",
            f"- Target center: `{_format_xy(target_station.get('center'))}`",
            f"- Target navigation approach: `{_format_xy(target_station.get('approach'))}`",
            "",
            "## BC Grasp Start Pose",
            "",
            "- No fixed BC XY coordinate is copied into this SOP.",
            "- The exact object pose must be read from the live MuJoCo scene at execution time.",
            "- If an object-specific grasp start pose is required, the manipulation skill must derive and validate it from live state.",
            "- Failure to resolve a live pose is an execution error; never substitute a remembered coordinate.",
        ]
    )

    lines.extend(
        [
            "",
            "## Required Skill Flow",
            "",
            f"1. `move(target=\"{source}\")`",
        ]
    )
    if selected_objects:
        if quantity == 1:
            lines.append(f"2. `pick_up(object_name=\"{selected_objects[0]}\")`")
        else:
            lines.append(
                "2. Pick the required objects in configured order: "
                + ", ".join(f"`{name}`" for name in selected_objects)
            )
    else:
        lines.append("2. `pick_up` — exact object name unresolved; do not guess")
    lines.extend(
        [
            f"3. `move(target=\"{target}\")` while carrying the object",
            f"4. `place_down(target=\"{target}\")`",
        ]
    )
    if quantity > 1:
        lines.append(
            f"5. Repeat the pick/transport/place cycle until `{quantity}` objects are placed."
        )

    lines.extend(
        [
            "",
            "## Execution Checks",
            "",
            "- Stop diagnosis at the first failed stage.",
            "- A successful move proves navigation reached its target; it does not prove the BC grasp pose is correct.",
            "- Do not treat a later place failure as independent when pick_up already failed.",
            "- If pick_up fails, verify exact base pose, yaw, object_name, online observations, and checkpoint contract before retraining.",
        ]
    )

    useful_descriptions = {
        name: _compact(description)[:1600]
        for name, description in image_descriptions.items()
        if description and not str(description).startswith("VLM error:")
    }
    lines.extend(
        [
            "",
            "## Document Image Analysis",
            "",
            f"- Embedded image count: `{image_count}`",
        ]
    )
    if useful_descriptions:
        lines.append(
            "- VLM descriptions are advisory visual context; they are never a coordinate authority."
        )
        for name, description in useful_descriptions.items():
            lines.extend([f"- `{name}`:", f"  {description}"])
    else:
        lines.append("- No successful VLM image descriptions were available during generation.")

    return "\n".join(lines).rstrip() + "\n"


class ReadDocumentSkill(BaseSkill):
    """Read a .docx, optionally describe images, and generate team SOP markdown.

    LLM can invoke this with::

        {"skill_name": "read_document",
         "inputs": {"file": "knowledge/JCIIOT_2026_case_1_SOP.docx",
                    "use_vision": true,
                    "generate_sop": true}}

    ``generate_sop`` and ``activate_sop`` default to ``True``.  The detailed
    result is archived under ``team_submission/generated_sops`` and copied to
    the single planner-visible
    ``team_submission/knowledge/current_generated_sop.md``.  Locked
    competition knowledge cannot be overwritten accidentally.
    """

    def __init__(
        self,
        *,
        ollama_base_url: str = "http://localhost:11434",
        vision_model: str = "qwen3-vl:8b",
        api_type: str = "ollama",
        api_key: str = "",
    ) -> None:
        super().__init__(
            name="read_document",
            description="Read .docx files, extract text and analyze images with vision model",
            keywords=("read", "document", "docx", "analyze", "vision", "parse", "extract"),
        )
        self._ollama_url = ollama_base_url
        self._vision_model = vision_model
        self._api_type = api_type
        self._api_key = api_key

    def run(self, context: ExecutionContext) -> SkillResult:
        inputs = context.metadata.get("inputs", {})
        file_path = inputs.get("file", "")
        use_vision = inputs.get("use_vision", True)
        generate_sop = inputs.get("generate_sop", True)
        activate_sop = inputs.get("activate_sop", True)

        path = Path(file_path)
        if not path.exists():
            return SkillResult(
                skill_name=self.name, success=False,
                message=f"File not found: {file_path}",
                payload={"file": file_path},
            )

        try:
            from docx import Document
            doc = Document(str(path))

            # Extract all paragraph text
            paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
            table_lines = []
            for table in doc.tables:
                for row in table.rows:
                    cells = [_compact(cell.text) for cell in row.cells]
                    if any(cells):
                        table_lines.append(" | ".join(cells))
            full_text = "\n".join(paragraphs + table_lines)

            # Extract images
            images = {}
            for rel in doc.part.rels.values():
                if "image" in rel.reltype:
                    name = rel.target_ref.split("/")[-1] if rel.target_ref else "image.png"
                    images[name] = rel.target_part.blob

            # Optional vision analysis
            img_descriptions = {}
            if use_vision and images:
                try:
                    from robot_agent.core.vision_client import ask_vision
                    for name, img_data in images.items():
                        try:
                            desc = ask_vision(
                                "Describe this factory layout image. What stations, tables, "
                                "production lines, objects, and their positions do you see?",
                                img_data,
                                base_url=self._ollama_url,
                                model=self._vision_model,
                                api_type=self._api_type,
                                api_key=self._api_key,
                            )
                            img_descriptions[name] = desc
                        except Exception as exc:
                            img_descriptions[name] = f"VLM error: {exc}"
                except Exception as exc:
                    logger.warning("Vision analysis skipped: %s", exc)

            generated_sop_path = ""
            active_sop_path = ""
            generated_sop_level = ""
            generated_sop = ""
            if generate_sop:
                prompt = _parse_prompt_fields(_extract_prompt_lines(paragraphs))
                scene_metadata = context.metadata.get("scene", {})
                if not isinstance(scene_metadata, dict):
                    scene_metadata = {}
                runtime_environment = (
                    inputs.get("environment")
                    or context.metadata.get("env_name")
                    or scene_metadata.get("env_name")
                    or ""
                )
                task, generated_sop_level, case_number = _resolve_task_entry(
                    path,
                    requested_level=str(inputs.get("level") or ""),
                    requested_environment=str(runtime_environment),
                )
                scene_map, map_path = _load_scene_map(task)
                generated_sop = _build_sop_markdown(
                    source_path=path,
                    level=generated_sop_level,
                    case_number=case_number,
                    prompt=prompt,
                    task=task,
                    scene_map=scene_map,
                    map_path=map_path,
                    image_count=len(images),
                    image_descriptions=img_descriptions,
                )
                output_path = _safe_output_path(
                    generated_sop_level,
                    case_number,
                    inputs,
                )
                output_path.write_text(generated_sop, encoding="utf-8")
                generated_sop_path = str(output_path)

                if activate_sop:
                    _TEAM_KNOWLEDGE_ROOT.mkdir(parents=True, exist_ok=True)
                    _ACTIVE_SOP_PATH.write_text(generated_sop, encoding="utf-8")
                    active_sop_path = str(_ACTIVE_SOP_PATH)

                    # Refresh the team index immediately. The planner also
                    # reloads this directory before building every prompt.
                    try:
                        from robot_agent.core.knowledge_manager import KnowledgeManager

                        manager = KnowledgeManager(_TEAM_KNOWLEDGE_ROOT)
                        manager.reload()
                        for document in manager.list_docs():
                            if document.get("source_file") == _ACTIVE_SOP_PATH.name:
                                manager.update_doc_meta(
                                    document["doc_id"],
                                    title=f"{generated_sop_level or 'Task'} Current Team SOP",
                                    category="sop",
                                    tags=[
                                        generated_sop_level or "task",
                                        "current",
                                        "team-generated",
                                        "SOP",
                                    ],
                                )
                                break
                    except Exception as exc:
                        logger.warning("Generated SOP index refresh skipped: %s", exc)

            return SkillResult(
                skill_name=self.name,
                success=True,
                message=f"Read {len(paragraphs)} paragraphs, {len(images)} images"
                        + (f", {len(img_descriptions)} analyzed by VLM" if img_descriptions else "")
                        + (f", generated {Path(generated_sop_path).name}" if generated_sop_path else ""),
                payload={
                    "file": str(path),
                    "paragraph_count": len(paragraphs),
                    "table_row_count": len(table_lines),
                    "image_count": len(images),
                    "images_analyzed": len(img_descriptions),
                    "text": full_text,
                    "image_descriptions": img_descriptions,
                    "generated_sop_path": generated_sop_path,
                    "active_sop_path": active_sop_path,
                    "generated_sop_level": generated_sop_level,
                    "generated_sop_markdown": generated_sop,
                },
            )
        except Exception as exc:
            logger.exception("ReadDocumentSkill failed")
            return SkillResult(
                skill_name=self.name, success=False,
                message=f"Failed: {exc}",
                payload={"file": file_path, "error": str(exc)},
            )
